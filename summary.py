pip install openai python-docx

import openai
from docx import Document
import os

# Set your OpenAI API key
openai.api_key = os.environ.get("OPENAI_API_KEY")  # Or replace with your API key string

def read_transcript(file_path):
    """Reads the transcript from a .txt file."""
    with open(file_path, 'r', encoding='utf-8') as file:
        transcript = file.read()
    return transcript

def summary_extraction(transcript):
    """Generates a summary of the transcript."""
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a highly skilled AI trained in language comprehension and summarization. "
                    "I would like you to read the following text and summarize it into a concise abstract paragraph. "
                    "Aim to retain the most important points, providing a coherent and readable summary that could help "
                    "a person understand the main points of the discussion without needing to read the entire text. "
                    "Please avoid unnecessary details or tangential points."
                )
            },
            {
                "role": "user",
                "content": transcript
            }
        ]
    )
    return response.choices[0].message.content.strip()

def key_points_extraction(transcript):
    """Extracts key points from the transcript."""
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are a proficient AI with a specialty in distilling information into key points. "
                    "Based on the following text, identify and list the main points that were discussed or brought up. "
                    "These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. "
                    "Your goal is to provide a list that someone could read to quickly understand what was talked about."
                )
            },
            {
                "role": "user",
                "content": transcript
            }
        ]
    )
    return response.choices[0].message.content.strip()

def action_items_extraction(transcript):
    """Identifies action items from the transcript."""
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "You are an AI expert in analyzing conversations and extracting action items. "
                    "Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. "
                    "These could be tasks assigned to specific individuals, or general actions that the group has decided to take. "
                    "Please list these action items clearly and concisely."
                )
            },
            {
                "role": "user",
                "content": transcript
            }
        ]
    )
    return response.choices[0].message.content.strip()

def sentiment_analysis(transcript):
    """Performs sentiment analysis on the transcript."""
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": (
                    "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. "
                    "Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. "
                    "Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
                )
            },
            {
                "role": "user",
                "content": transcript
            }
        ]
    )
    return response.choices[0].message.content.strip()

def generate_meeting_minutes(transcript):
    """Generates the meeting minutes by calling the extraction functions."""
    summary = summary_extraction(transcript)
    key_points = key_points_extraction(transcript)
    action_items = action_items_extraction(transcript)
    sentiment = sentiment_analysis(transcript)
    return {
        'Summary': summary,
        'Key Points': key_points,
        'Action Items': action_items,
        'Sentiment Analysis': sentiment
    }

def save_as_docx(minutes, filename):
    """Saves the meeting minutes into a Word document."""
    doc = Document()
    for section, content in minutes.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content)
        doc.add_paragraph()
    doc.save(filename)

if __name__ == "__main__":
    # Replace 'transcript.txt' with the path to the transcript file
    transcript_file_path = "transcript.txt"
    transcript = read_transcript(transcript_file_path)
    minutes = generate_meeting_minutes(transcript)
    
    # Print the results to the console
    for section, content in minutes.items():
        print(f"--- {section} ---\n{content}\n")

    # Save the results to a Word document
    save_as_docx(minutes, 'meeting_minutes.docx')
    print("Meeting minutes have been saved to 'meeting_minutes.docx'.")


